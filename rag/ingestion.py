"""Document ingestion: validate and extract text from PDF/DOCX/TXT."""
import os
from dataclasses import dataclass
from typing import List, Tuple
from .logging_utils import get_logger, Timer
from .config import RAGConfig

logger = get_logger("ingestion")


@dataclass
class IngestedDocument:
    """Container for an ingested document with per-page text."""
    source_path: str
    source_filename: str
    pages_text: List[Tuple[int, str]]
    total_pages: int


def _extract_pdf(path: str, cfg: RAGConfig) -> IngestedDocument:
    import fitz  # PyMuPDF
    doc = fitz.open(path)
    total_pages = doc.page_count
    if total_pages > cfg.max_pages_pdf:
        logger.info(
            "PDF exceeds page limit",
            extra={"event": "page_count_check", "data": {"pages": total_pages}},
        )
        raise ValueError(f"PDF has {total_pages} pages; limit is {cfg.max_pages_pdf}")
    pages_text: List[Tuple[int, str]] = []
    for i in range(total_pages):
        page = doc.load_page(i)
        text = page.get_text("text")
        pages_text.append((i + 1, text or ""))
    doc.close()
    return IngestedDocument(path, os.path.basename(path), pages_text, total_pages)


def _extract_docx(path: str) -> IngestedDocument:
    import docx
    d = docx.Document(path)
    text = "\n".join(p.text for p in d.paragraphs)
    pages_text = [(1, text)]
    return IngestedDocument(path, os.path.basename(path), pages_text, 1)


def _extract_txt(path: str) -> IngestedDocument:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return IngestedDocument(path, os.path.basename(path), [(1, text)], 1)


def ingest_document(path: str, cfg: RAGConfig) -> IngestedDocument:
    """Extract text per page from the given file and enforce constraints."""
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")
    ext = os.path.splitext(path.lower())[1]
    timer = Timer()
    logger.info(
        "File upload",
        extra={"event": "file_upload", "data": {"file": os.path.basename(path), "ext": ext}},
    )
    if ext == ".pdf":
        ingested = _extract_pdf(path, cfg)
    elif ext == ".docx":
        ingested = _extract_docx(path)
    elif ext == ".txt":
        ingested = _extract_txt(path)
    else:
        raise ValueError("Unsupported file type. Allowed: PDF, DOCX, TXT")
    logger.info(
        "Extraction complete",
        extra={
            "event": "extraction",
            "data": {"pages": ingested.total_pages, "ms": timer.elapsed_ms()},
        },
    )
    return ingested
